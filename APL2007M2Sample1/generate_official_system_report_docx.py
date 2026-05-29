from __future__ import annotations

import base64
from datetime import date
from pathlib import Path

import requests
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_NAME = "APL2007M2Sample1"
REPORT_TITLE = f"{PROJECT_NAME} 系統分析與設計報告書"
VERSION = "v1.0"
DOC_ID = "SYS-APL2007M2-001"
DATE_STR = date.today().isoformat()
OUTPUT_DOCX = "System_Report_APL2007M2Sample1_Official.docx"
ASSET_DIR = Path("report_assets")

ARCH_MERMAID = """flowchart LR
    U[使用者] --> B[Start 按鈕]
    B --> E[OnStartButtonClick]
    E --> BG[Task.Run 背景執行]
    BG --> S[StartSumPageSizesAsync]
    S --> P[SumPageSizesAsync]
    P --> W[Task.WhenAll 平行下載]
    W --> H[HttpClient GET 請求]
    H --> R[逐筆結果寫入 TextBox]
    W --> T[彙總總位元組與耗時]
    T --> UI[Dispatcher 切回 UI 執行緒]
    UI --> D[顯示結果並重新啟用按鈕]
"""

SEQ_MERMAID = """sequenceDiagram
    participant User as 使用者
    participant UI as MainWindow (UI Thread)
    participant BG as Background Task
    participant HTTP as HttpClient
    participant DISP as Dispatcher

    User->>UI: 點擊 Start
    UI->>UI: 停用按鈕並清空結果
    UI->>BG: Task.Run(StartSumPageSizesAsync)
    BG->>BG: 為每個 URL 建立 Task
    BG->>HTTP: 平行送出 GET 請求

    loop 每個 URL 完成時
        HTTP-->>BG: 回傳 byte[]
        BG->>DISP: BeginInvoke(附加單筆結果)
        DISP->>UI: 更新 TextBox
    end

    BG->>BG: Await Task.WhenAll 並加總
    BG->>DISP: BeginInvoke(附加總量與耗時)
    DISP->>UI: 啟用 Start 按鈕
"""


def set_default_font(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("頁碼：")
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "請在 Word 中按 F9 更新目錄"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.extend([fld_char_begin, instr_text, fld_char_separate, text, fld_char_end])


def mermaid_to_png(mermaid_text: str, output_path: Path) -> None:
    encoded = base64.urlsafe_b64encode(mermaid_text.encode("utf-8")).decode("ascii")
    url = f"https://mermaid.ink/img/{encoded}?type=png"
    response = requests.get(url, timeout=30)
    response.raise_for_status()
    output_path.write_bytes(response.content)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.runs[0].font.name = "Microsoft JhengHei"
    p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def build_report() -> None:
    ASSET_DIR.mkdir(exist_ok=True)
    arch_png = ASSET_DIR / "architecture.png"
    seq_png = ASSET_DIR / "sequence.png"

    mermaid_to_png(ARCH_MERMAID, arch_png)
    mermaid_to_png(SEQ_MERMAID, seq_png)

    doc = Document()
    set_default_font(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(REPORT_TITLE)
    run.bold = True
    run.font.size = Pt(22)

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("公司正式文件").bold = True

    for line in [
        f"文件編號：{DOC_ID}",
        f"版本：{VERSION}",
        f"日期：{DATE_STR}",
        "部門：資訊部",
        "作者：系統分析人員",
        "公司：＿＿＿＿＿＿＿＿＿＿",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    add_heading(doc, "文件管制資訊", 1)
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    info_rows = [
        ("文件名稱", REPORT_TITLE),
        ("文件編號", DOC_ID),
        ("版本", VERSION),
        ("建立日期", DATE_STR),
        ("適用系統", PROJECT_NAME),
        ("文件等級", "內部使用"),
    ]
    for idx, (k, v) in enumerate(info_rows):
        table.cell(idx, 0).text = k
        table.cell(idx, 1).text = v

    add_heading(doc, "修訂紀錄", 1)
    rev = doc.add_table(rows=2, cols=5)
    rev.style = "Table Grid"
    rev.cell(0, 0).text = "版本"
    rev.cell(0, 1).text = "日期"
    rev.cell(0, 2).text = "修訂內容"
    rev.cell(0, 3).text = "作者"
    rev.cell(0, 4).text = "備註"
    rev.cell(1, 0).text = VERSION
    rev.cell(1, 1).text = DATE_STR
    rev.cell(1, 2).text = "初版建立"
    rev.cell(1, 3).text = "Copilot"
    rev.cell(1, 4).text = "-"

    add_heading(doc, "簽核欄", 1)
    sign = doc.add_table(rows=2, cols=3)
    sign.style = "Table Grid"
    sign.cell(0, 0).text = "角色"
    sign.cell(0, 1).text = "姓名"
    sign.cell(0, 2).text = "簽核"
    sign.cell(1, 0).text = "部門主管"
    sign.cell(1, 1).text = ""
    sign.cell(1, 2).text = ""

    doc.add_page_break()

    add_heading(doc, "目錄", 1)
    toc_p = doc.add_paragraph()
    add_toc(toc_p)

    doc.add_page_break()

    add_heading(doc, "1. 專案概述", 1)
    doc.add_paragraph(
        "APL2007M2Sample1 為 .NET 6 WPF 桌面應用範例，展示平行非同步下載與 UI 執行緒切換。"
        "使用者點擊 Start 後，系統會同時向多個 Microsoft Docs URL 發送請求，逐筆輸出下載大小，"
        "並於結尾顯示總位元組數與總耗時。"
    )

    add_heading(doc, "2. 專案目標", 1)
    goals = [
        "展示 async/await 與 Task.WhenAll 的平行處理實務。",
        "展示背景工作與 WPF UI 更新協調方式。",
        "展示多任務結果彙整與輸出。",
    ]
    for g in goals:
        doc.add_paragraph(g, style="List Number")

    add_heading(doc, "3. 技術架構", 1)
    bullets = [
        "語言：C#",
        "框架：.NET 6（net6.0-windows）",
        "UI：WPF",
        "HTTP：System.Net.Http.HttpClient",
        "非同步：Task / async/await / Task.WhenAll",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    add_heading(doc, "4. 系統架構圖", 1)
    doc.add_picture(str(arch_png), width=Cm(16))
    cap1 = doc.add_paragraph("圖 4-1 系統架構圖")
    cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "5. 執行時序圖", 1)
    doc.add_picture(str(seq_png), width=Cm(16))
    cap2 = doc.add_paragraph("圖 5-1 執行時序圖")
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "6. 模組與檔案職責", 1)
    modules = [
        "APL2007M2Sample1.csproj：WPF 專案設定與目標框架。",
        "App.xaml：啟動頁定義（StartupUri=MainWindow.xaml）。",
        "App.xaml.cs：Application 類別。",
        "MainWindow.xaml：UI 元件（Start 按鈕與結果 TextBox）。",
        "MainWindow.xaml.cs：下載流程、結果彙整、Dispatcher UI 更新。",
    ]
    for m in modules:
        doc.add_paragraph(m, style="List Bullet")

    add_heading(doc, "7. 功能規格摘要", 1)
    features = [
        "一鍵觸發 19 個 URL 平行下載。",
        "逐筆顯示每個 URL 的位元組數。",
        "顯示總位元組與耗時。",
        "執行中停用按鈕避免重複觸發。",
        "關閉視窗時釋放 HttpClient。",
    ]
    for f in features:
        doc.add_paragraph(f, style="List Number")

    add_heading(doc, "8. 執行環境與部署方式", 1)
    doc.add_paragraph("前置需求：Windows、.NET 6 SDK。")
    doc.add_paragraph("建置與執行指令：")
    code = doc.add_paragraph("dotnet restore\ndotnet build\ndotnet run")
    for run in code.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(10)

    add_heading(doc, "9. 測試現況", 1)
    doc.add_paragraph("目前本資料夾無獨立測試專案，主要透過手動執行介面進行驗證。")

    add_heading(doc, "10. 風險評估", 1)
    risks = [
        "Task.WhenAll 缺少完整例外處理，任一請求失敗可能中斷流程。",
        "Task.Run fire-and-forget 型態使例外可觀測性較低。",
        "缺少取消機制與重試退避策略。",
        "URL 清單硬編碼，不利於營運期調整。",
        "邏輯集中於單一視窗類別，擴充可維護性較弱。",
    ]
    for r in risks:
        doc.add_paragraph(r, style="List Number")

    add_heading(doc, "11. 後續建議", 1)
    suggestions = [
        "加入 try/catch/finally 並確保 UI 狀態復原。",
        "加入 CancellationToken 與 Cancel 按鈕。",
        "URL 外部化設定（如 appsettings.json）。",
        "加入結構化日誌與監控。",
        "下載邏輯服務化並補齊單元測試。",
    ]
    for s in suggestions:
        doc.add_paragraph(s, style="List Number")

    add_heading(doc, "12. 結論", 1)
    doc.add_paragraph(
        "本系統具備明確教學價值，可快速展示桌面應用中的平行非同步模式。"
        "若要進一步產品化，建議優先補強錯誤處理、可取消控制、設定外部化與自動化測試。"
    )

    doc.save(OUTPUT_DOCX)
    print(f"已產生：{OUTPUT_DOCX}")
    print(f"圖檔輸出：{arch_png}、{seq_png}")


if __name__ == "__main__":
    build_report()
