from datetime import date
from docx import Document
from docx.shared import Pt


def add_code_block(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def main() -> None:
    doc = Document()

    doc.add_heading("APL2007M2Sample1 系統報告書", level=0)
    doc.add_paragraph(f"文件日期：{date.today().isoformat()}")
    doc.add_paragraph("文件版本：v1.0")
    doc.add_paragraph("文件狀態：初版")

    doc.add_heading("1. 文件摘要", level=1)
    doc.add_paragraph(
        "本文件針對 APL2007M2Sample1 專案進行系統化整理，包含專案目的、技術架構、執行流程、"
        "功能內容、部署方式、風險限制與後續建議，提供接手維護與持續演進所需之基礎資訊。"
    )

    doc.add_heading("2. 專案基本資訊", level=1)
    add_bullets(
        doc,
        [
            "專案名稱：APL2007M2Sample1",
            "專案型態：Windows 桌面應用程式（WPF）",
            "目標框架：.NET 6（net6.0-windows）",
            "主要語言：C#",
            "啟動方式：由 App.xaml 指定 StartupUri 為 MainWindow.xaml",
        ],
    )

    doc.add_heading("3. 專案目的與範圍", level=1)
    doc.add_paragraph("本專案為教學範例，目的在於展示桌面應用中的平行非同步處理與 UI 執行緒切換。")
    add_numbers(
        doc,
        [
            "示範如何以 async/await 與 Task.WhenAll 進行平行下載。",
            "示範如何在背景工作執行時避免 UI 卡死。",
            "示範如何安全地將背景執行結果回寫至 WPF UI。",
        ],
    )

    doc.add_heading("4. 技術架構", level=1)
    add_bullets(
        doc,
        [
            "UI：WPF（MainWindow）",
            "非同步模型：Task、async/await、Task.WhenAll",
            "網路通訊：System.Net.Http.HttpClient",
            "執行緒切換：Dispatcher.BeginInvoke",
        ],
    )

    doc.add_paragraph("系統架構圖（Mermaid 原始碼）：")
    add_code_block(
        doc,
        """```mermaid
flowchart LR
    U[User] --> B[Start Button]
    B --> E[OnStartButtonClick]
    E --> BG[Task.Run background operation]
    BG --> S[StartSumPageSizesAsync]
    S --> P[SumPageSizesAsync]
    P --> W[Task.WhenAll parallel downloads]
    W --> H[HttpClient GET requests]
    H --> R[Per-URL result appended to TextBox]
    W --> T[Aggregate total bytes and elapsed time]
    T --> UI[Dispatcher marshals to UI thread]
    UI --> D[Render summary and re-enable button]
```""",
    )

    doc.add_heading("5. 主要流程設計", level=1)
    add_numbers(
        doc,
        [
            "使用者點擊 Start 按鈕。",
            "系統停用按鈕並清空結果區。",
            "建立多個下載任務並以 Task.WhenAll 平行執行。",
            "每個任務完成後將 URL 與位元組大小追加到結果區。",
            "全部完成後彙總總位元組與耗時並顯示。",
            "恢復 Start 按鈕可點擊狀態。",
        ],
    )

    doc.add_paragraph("執行時序圖（Mermaid 原始碼）：")
    add_code_block(
        doc,
        """```mermaid
sequenceDiagram
    participant User as User
    participant UI as MainWindow (UI Thread)
    participant BG as Background Task
    participant HTTP as HttpClient
    participant DISP as Dispatcher

    User->>UI: Click Start
    UI->>UI: Disable Start and clear results
    UI->>BG: Task.Run(StartSumPageSizesAsync)
    BG->>BG: Build one task per URL
    BG->>HTTP: Send parallel GET requests

    loop For each URL completion
        HTTP-->>BG: Return byte[]
        BG->>DISP: BeginInvoke append per-URL size
        DISP->>UI: Update TextBox
    end

    BG->>BG: Await Task.WhenAll and sum lengths
    BG->>DISP: BeginInvoke append total and elapsed time
    DISP->>UI: Re-enable Start button
```""",
    )

    doc.add_heading("6. 模組與檔案說明", level=1)
    add_bullets(
        doc,
        [
            "APL2007M2Sample1.csproj：宣告 WPF、WinExe、目標框架。",
            "App.xaml：應用程式啟動入口設定。",
            "App.xaml.cs：Application 類別。",
            "MainWindow.xaml：主視窗 UI 定義（Start 按鈕與結果文字框）。",
            "MainWindow.xaml.cs：核心邏輯（下載、彙總、UI 更新、資源釋放）。",
        ],
    )

    doc.add_heading("7. 功能內容", level=1)
    add_bullets(
        doc,
        [
            "一鍵觸發 19 個 URL 的平行下載。",
            "逐筆輸出每個 URL 的下載位元組數。",
            "輸出總位元組數與執行耗時。",
            "執行期間禁止重複點擊 Start。",
            "視窗關閉時釋放 HttpClient。",
        ],
    )

    doc.add_heading("8. 執行環境與部署", level=1)
    add_bullets(
        doc,
        [
            "作業系統：Windows",
            "必要元件：.NET 6 SDK",
            "建置指令：dotnet restore、dotnet build",
            "執行指令：dotnet run",
        ],
    )

    doc.add_heading("9. 驗證與測試現況", level=1)
    doc.add_paragraph("目前專案資料夾內未包含獨立測試專案，主要以手動執行驗證功能。")

    doc.add_heading("10. 風險與限制", level=1)
    add_numbers(
        doc,
        [
            "缺少 Task.WhenAll 外層例外處理，任一下載失敗可能導致整體流程拋例外。",
            "Click handler 以 Task.Run 啟動背景工作，屬 fire-and-forget 型態，例外可觀測性較弱。",
            "目前沒有取消機制與逾時控制策略。",
            "URL 清單硬編碼於程式中，缺少外部化設定。",
            "邏輯集中於單一視窗類別，擴充時可維護性較弱。",
        ],
    )

    doc.add_heading("11. 改進建議", level=1)
    add_numbers(
        doc,
        [
            "加入 try/catch/finally，確保異常時仍能恢復 UI 狀態。",
            "加入 CancellationToken 與取消按鈕。",
            "將 URL 清單移至設定檔。",
            "導入重試與退避策略，提升網路穩定性。",
            "抽離下載服務層並補齊單元測試。",
        ],
    )

    doc.add_heading("12. 維運交接重點", level=1)
    add_bullets(
        doc,
        [
            "核心邏輯集中於 MainWindow.xaml.cs，可優先閱讀此檔理解全流程。",
            "目前主要依賴外部網路狀態，建議先確認網路環境再排查程式問題。",
            "若要產品化，應優先處理例外、取消、設定外部化、測試自動化。",
        ],
    )

    doc.add_heading("13. 參考來源", level=1)
    add_bullets(
        doc,
        [
            "README.md",
            "APL2007M2Sample1.csproj",
            "App.xaml / App.xaml.cs",
            "MainWindow.xaml / MainWindow.xaml.cs",
        ],
    )

    output_path = "System_Report_APL2007M2Sample1.docx"
    doc.save(output_path)
    print(f"已產生：{output_path}")


if __name__ == "__main__":
    main()
