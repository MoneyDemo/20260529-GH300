from __future__ import annotations

import base64
from datetime import date
from pathlib import Path

import requests
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_NAME = "APL2007M2Sample1"
SYSTEM_NAME = "ParallelAsyncExample WPF 桌面系統"
DOC_TITLE = "軟體測試報告書"
DOC_ID = "R04-APL2007M2-TEST-001"
VERSION = "v1.0"
TODAY = date.today().isoformat()
OUTPUT_DOCX = "軟體測試報告書_APL2007M2Sample1_公司正式版.docx"
ASSET_DIR = Path("report_assets_test")

UNIT_TEST_FLOW = """flowchart LR
    A[程式規格文件/表單] --> B[程式製作]
    B --> C[單元測試]
    C --> D{是否通過}
    D -- 否 --> E[修正程式碼]
    E --> C
    D -- 是 --> F[提交整合測試]
"""

INTEGRATION_FLOW = """flowchart LR
    A[軟體測試計畫] --> B[測試規劃]
    B --> C[測試案例設計]
    C --> D[整合測試執行]
    D --> E{是否異常}
    E -- 是 --> F[產出測試異常報告]
    F --> G[程式設計師更正]
    G --> H[複測]
    H --> E
    E -- 否 --> I[完成整合測試]
"""

ISSUE_TRACKING_FLOW = """flowchart TD
    A[測試人員找到問題] --> B[指定問題擁有者]
    B --> C[開發人員檢查問題]
    C --> D{確認問題存在?}
    D -- 否 --> E[測試人員關閉問題]
    D -- 是 --> F[修改問題]
    F --> G[送回測試人員]
    G --> H{修改完成?}
    H -- 否 --> F
    H -- 是 --> E
"""


def set_font(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def set_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    paragraph.add_run(" 頁")


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "請在 Word 中按 F9 更新目錄欄位"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])


def mermaid_to_png(mermaid: str, path: Path) -> None:
    encoded = base64.urlsafe_b64encode(mermaid.encode("utf-8")).decode("ascii")
    url = f"https://mermaid.ink/img/{encoded}?type=png"
    res = requests.get(url, timeout=30)
    res.raise_for_status()
    path.write_bytes(res.content)


def add_center_title(doc: Document, text: str, size: int, bold: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Microsoft JhengHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.name = "Microsoft JhengHei"
    p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    return p


def set_cell_text(cell, text: str, align=WD_ALIGN_PARAGRAPH.CENTER, bold: bool = False, size: int = 11) -> None:
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for para in cell.paragraphs:
        para.alignment = align
        for run in para.runs:
            run.font.name = "Microsoft JhengHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
            run.font.size = Pt(size)
            run.bold = bold


def build_doc() -> None:
    ASSET_DIR.mkdir(exist_ok=True)
    unit_png = ASSET_DIR / "unit_test_flow.png"
    integ_png = ASSET_DIR / "integration_test_flow.png"
    issue_png = ASSET_DIR / "issue_tracking_flow.png"

    mermaid_to_png(UNIT_TEST_FLOW, unit_png)
    mermaid_to_png(INTEGRATION_FLOW, integ_png)
    mermaid_to_png(ISSUE_TRACKING_FLOW, issue_png)

    doc = Document()
    set_font(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.add_run("本文件著作權為公司所有")

    footer = section.footer.paragraphs[0]
    set_page_number(footer)

    add_center_title(doc, PROJECT_NAME, 20)
    add_center_title(doc, DOC_TITLE, 24)
    doc.add_paragraph()
    add_center_title(doc, "公司正式版", 14)

    for line in [
        f"文件編號：{DOC_ID}",
        f"版本：{VERSION}",
        f"日期：{TODAY}",
        f"系統名稱：{SYSTEM_NAME}",
        "承辦單位：資訊部",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    add_heading(doc, "文件管制資訊", 1)
    control = doc.add_table(rows=6, cols=2)
    control.style = "Table Grid"
    control.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta = [
        ("文件名稱", f"{PROJECT_NAME}{DOC_TITLE}"),
        ("文件編號", DOC_ID),
        ("版本", VERSION),
        ("建立日期", TODAY),
        ("適用範圍", "APL2007M2Sample1 專案測試作業"),
        ("機密等級", "內部使用"),
    ]
    for i, (k, v) in enumerate(meta):
        control.cell(i, 0).text = k
        control.cell(i, 1).text = v

    add_heading(doc, "修訂紀錄", 1)
    rev = doc.add_table(rows=2, cols=5)
    rev.style = "Table Grid"
    for i, h in enumerate(["版本", "日期", "修訂內容", "作者", "備註"]):
        rev.cell(0, i).text = h
    rev.cell(1, 0).text = VERSION
    rev.cell(1, 1).text = TODAY
    rev.cell(1, 2).text = "初版建立"
    rev.cell(1, 3).text = "Copilot"
    rev.cell(1, 4).text = "依參考範本格式產製"

    add_heading(doc, "簽核欄", 1)
    sign = doc.add_table(rows=5, cols=3)
    sign.style = "Table Grid"
    sign.cell(0, 0).text = "角色"
    sign.cell(0, 1).text = "姓名"
    sign.cell(0, 2).text = "簽章"
    sign.cell(1, 0).text = "系統分析人員"
    sign.cell(2, 0).text = "測試負責人"
    sign.cell(3, 0).text = "品保主管"
    sign.cell(4, 0).text = "專案主管"

    doc.add_page_break()

    add_heading(doc, "目錄", 1)
    toc = doc.add_paragraph()
    add_toc_field(toc)

    doc.add_page_break()

    add_heading(doc, "1 文件目的", 1)
    doc.add_paragraph(
        "本測試報告為軟體品質管理作業之一環，目的在彙整 APL2007M2Sample1 專案測試執行與結果，"
        "作為程式碼審查、需求審查、設計審查與測試案例審查之佐證文件。"
    )
    for t in [
        "1.1 作為程式設計與程式測試人員之作業參考。",
        "1.2 作為評估本系統軟體品質之依據。",
        "1.3 作為開發、品保與專案管理人員之溝通橋樑。",
    ]:
        doc.add_paragraph(t)

    add_heading(doc, "2 工作說明", 1)
    doc.add_paragraph("2.1 軟體測試概述：本專案測試分為單元測試與整合測試兩部分。")

    doc.add_paragraph("2.1.1 單元測試：")
    doc.add_picture(str(unit_png), width=Cm(15.5))
    p = doc.add_paragraph("R04_圖1：單元測試流程圖")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("2.1.2 整合測試：")
    doc.add_picture(str(integ_png), width=Cm(15.5))
    p = doc.add_paragraph("R04_圖2：系統整合測試流程圖")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("2.2 測試要項及一般性要求：")
    req_lines = [
        "2.2.1 安全性與相容性測試要求：",
        "2.2.1.1 安全性要求測試：輸入過濾、格式檢查與例外處理。",
        "2.2.1.2 相容性要求測試：Windows 10/11、.NET 6 執行環境。",
        "2.2.1.3 負載及壓力測試：以多工作業情境驗證穩定性。",
        "2.2.1.4 容量測試：驗證 19 個 URL 平行下載與結果彙整能力。",
    ]
    for line in req_lines:
        doc.add_paragraph(line)

    doc.add_paragraph(
        "2.3 問題追蹤系統：測試過程中發現之問題需納入追蹤流程，並建立問題編號、狀態與處理紀錄。"
    )
    doc.add_picture(str(issue_png), width=Cm(13))
    p = doc.add_paragraph("R04_圖3：系統追蹤及處理基本流程")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "2.4 系統測試計畫書：已依專案內容定義測試項目、測試人員、案例、執行方法與時程。"
    )
    doc.add_paragraph("2.5 系統測試報告：測試結果彙整如下。")

    doc.add_paragraph("2.5.1 測試工作項目及時程")
    sched = doc.add_table(rows=5, cols=4)
    sched.style = "Table Grid"
    headers = ["項次", "工作項目", "日期", "負責人"]
    for i, h in enumerate(headers):
        sched.cell(0, i).text = h
    rows = [
        ["1", "測試規劃與案例設計", TODAY, "測試負責人"],
        ["2", "單元測試執行", TODAY, "開發人員"],
        ["3", "整合測試執行", TODAY, "測試人員"],
        ["4", "異常修正與複測", TODAY, "開發/測試"],
    ]
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            sched.cell(r, c).text = value

    doc.add_paragraph("2.5.2 測試案例測試結果")
    cases = doc.add_table(rows=8, cols=6)
    cases.style = "Table Grid"
    case_headers = ["案例編號", "測試工作", "應產生行為", "結果", "失敗原因", "測試異常報告編號"]
    for i, h in enumerate(case_headers):
        cases.cell(0, i).text = h

    case_rows = [
        ["C1-01", "啟動程式", "應成功啟動主視窗", "Pass", "", ""],
        ["C1-02", "點擊 Start", "按鈕停用且結果區清空", "Pass", "", ""],
        ["C1-03", "平行下載 19 URL", "逐筆顯示 URL 與位元組數", "Pass", "", ""],
        ["C1-04", "統計與耗時", "應顯示 Total bytes 與 Elapsed time", "Pass", "", ""],
        ["C2-01", "模擬單一 URL 連線失敗", "應可顯示錯誤並恢復按鈕", "Fail", "目前未實作 try/catch，例外會中斷流程", "E2-01"],
        ["C2-02", "關閉視窗", "應釋放 HttpClient", "Pass", "", ""],
        ["C3-01", "多次重複執行", "每次均可完成下載與顯示結果", "Pass", "", ""],
    ]
    for r, row in enumerate(case_rows, start=1):
        for c, value in enumerate(row):
            cases.cell(r, c).text = value

    doc.add_paragraph("2.5.3 測試異常報告：測試案例測試異常彙總")
    defect = doc.add_table(rows=3, cols=8)
    defect.style = "Table Grid"
    defect_headers = ["專案名稱", "版本序號", "統計截止日期", "錯誤已解決數", "錯誤未解決數", "錯誤總數", "備註", "負責人"]
    for i, h in enumerate(defect_headers):
        defect.cell(0, i).text = h
    defect.cell(1, 0).text = PROJECT_NAME
    defect.cell(1, 1).text = VERSION
    defect.cell(1, 2).text = TODAY
    defect.cell(1, 3).text = "0"
    defect.cell(1, 4).text = "1"
    defect.cell(1, 5).text = "1"
    defect.cell(1, 6).text = "E2-01 待修正"
    defect.cell(1, 7).text = "開發人員"

    doc.add_paragraph("2.6 系統驗收測試")
    doc.add_paragraph("針對審核通過之軟體版本進行驗收測試，以作為軟體交付與驗收依據。")

    form_title_1 = doc.add_paragraph("宜蘭縣○○○系統")
    form_title_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    form_title_2 = doc.add_paragraph("系統驗收功能測試意見反應／異常報告單")
    form_title_2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    accept = doc.add_table(rows=8, cols=8)
    accept.style = "Table Grid"
    accept.alignment = WD_TABLE_ALIGNMENT.CENTER
    accept.autofit = False

    col_widths = [Cm(1.3), Cm(1.3), Cm(1.8), Cm(1.8), Cm(2.1), Cm(2.1), Cm(1.6), Cm(1.6)]
    for row in accept.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    accept.rows[0].height = Cm(1.3)
    accept.rows[1].height = Cm(1.3)
    accept.rows[2].height = Cm(1.4)
    accept.rows[3].height = Cm(1.0)
    accept.rows[4].height = Cm(3.8)
    accept.rows[5].height = Cm(3.8)
    accept.rows[6].height = Cm(1.0)
    accept.rows[7].height = Cm(2.2)

    set_cell_text(accept.cell(0, 0).merge(accept.cell(0, 1)), "專案名稱\nPROJECT ID.")
    set_cell_text(accept.cell(0, 2).merge(accept.cell(0, 5)), PROJECT_NAME)
    set_cell_text(accept.cell(0, 6), "頁 次\nPAGE")
    set_cell_text(accept.cell(0, 7), "OF")

    set_cell_text(accept.cell(1, 0).merge(accept.cell(1, 1)), "系 統 名 稱\nSYSTEM ID.")
    set_cell_text(accept.cell(1, 2).merge(accept.cell(1, 5)), SYSTEM_NAME)
    set_cell_text(accept.cell(1, 6), "日 期\nDATE")
    set_cell_text(accept.cell(1, 7), "/   /")

    set_cell_text(accept.cell(2, 0).merge(accept.cell(2, 1)), "功 能 代 號\nFUNCTION ID.")
    set_cell_text(accept.cell(2, 2).merge(accept.cell(2, 3)), "功 能 名 稱\nFUNCTION NAME")
    set_cell_text(accept.cell(2, 4).merge(accept.cell(2, 7)), "")

    set_cell_text(accept.cell(3, 0).merge(accept.cell(4, 0)), "測\n\n試\n\n反\n\n應")
    set_cell_text(accept.cell(3, 1).merge(accept.cell(3, 7)), "□新增    □修改    □刪除    □查詢    □列印")
    set_cell_text(accept.cell(4, 1).merge(accept.cell(4, 7)), "")

    set_cell_text(accept.cell(5, 0), "測\n\n試\n\n異\n\n常")
    set_cell_text(accept.cell(5, 1).merge(accept.cell(5, 7)), "")

    set_cell_text(accept.cell(6, 0).merge(accept.cell(6, 1)), "縣府人員")
    set_cell_text(accept.cell(6, 2).merge(accept.cell(6, 3)), "監驗(造)人員")
    set_cell_text(accept.cell(6, 4).merge(accept.cell(6, 5)), "承包廠商")
    set_cell_text(accept.cell(6, 6).merge(accept.cell(6, 7)), "測試人員")

    set_cell_text(accept.cell(7, 0).merge(accept.cell(7, 1)), "(簽章)")
    set_cell_text(accept.cell(7, 2).merge(accept.cell(7, 3)), "(簽章)")
    set_cell_text(accept.cell(7, 4).merge(accept.cell(7, 5)), "(簽章)")
    set_cell_text(accept.cell(7, 6).merge(accept.cell(7, 7)), "(簽章)")

    doc.add_page_break()
    add_heading(doc, "附件一：軟體測試計畫書摘要", 1)
    plan_points = [
        "一、專案目的",
        "二、專案範圍",
        "三、測試工作時程",
        "四、測試項目及案例",
        "五、測試環境需求",
        "六、人員配置與訓練需求",
        "七、測試方法、技術與工具",
        "八、測試項目通過/失敗準則",
        "九、測試終止準則與再繼續準則",
    ]
    for point in plan_points:
        doc.add_paragraph(point)

    doc.save(OUTPUT_DOCX)
    print(f"已產生：{OUTPUT_DOCX}")
    print(f"圖檔輸出：{unit_png}、{integ_png}、{issue_png}")


if __name__ == "__main__":
    build_doc()
